import os
import sys

import google.generativeai as genai
import instructor
from pydantic import BaseModel, Field
from haystack import component
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

from model.report_llm import report_model_response 

class Report_class(BaseModel):
    img_desc: str = Field(..., title="Description of the image")
    header: str
    img: str = Field(..., title="Path to the image")
    content: str

@component
class ReportGenerator:
    
    def create_word_report(report_data: Report_class, output_path: str):
        # Create a new Word document
        doc = Document()

        # Add a title (optional)
        title = doc.add_heading(level=1)
        title_run = title.add_run("Generated Report")
        title_run.bold = True
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add the header in bold
        header = doc.add_heading(level=2)
        header_run = header.add_run(report_data['header'])
        header_run.bold = True

        # Add the image description in italics
        img_desc_paragraph = doc.add_paragraph()
        img_desc_run = img_desc_paragraph.add_run(report_data['img_desc'])
        img_desc_run.italic = True

        # Add the image if provided
        if report_data['img']:
            doc.add_picture(report_data['img'], width=docx.shared.Inches(6))

        # Add the main content
        content_paragraph = doc.add_paragraph(report_data['content'])

        # Save the document
        doc.save(output_path)

    @component.output_types(report=str)
    def run(self, db_output: str):
        response_check = report_model_response(
            response_model_class=List[Report_class],
            report_prompt=create_report_prompt(db_output=db_output),
        )




